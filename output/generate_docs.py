#!/usr/bin/env python3
"""Generate professional Word documents for CM2121 submission."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)


def add_toc(doc):
    """Insert a Table of Contents field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
    )
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)
    run4 = p.add_run("(Right-click and select 'Update Field' to generate table of contents)")
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.size = Pt(11)
    run5 = p.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)


def create_cover_page(doc, title, subtitle="", extra_lines=None):
    """Create a professional cover page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CM2121 — 3D Reconstructive Techniques")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    if subtitle:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(80, 80, 80)

    for _ in range(4):
        doc.add_paragraph()

    lines = extra_lines or [
        ("Student:", "Matthew Jacob SD"),
        ("Student ID:", "[ID Placeholder]"),
        ("University:", "University of [University Name]"),
        ("Date:", "30 July 2026"),
        ("Extension Deadline:", "6 August 2026"),
    ]
    for label, value in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}  ")
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(value)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(60, 60, 60)

    doc.add_page_break()


def apply_styles(doc):
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Calibri"
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

    hs1 = doc.styles["Heading 1"]
    hs1.font.size = Pt(18)
    hs1.font.bold = True
    hs1.paragraph_format.space_before = Pt(24)

    hs2 = doc.styles["Heading 2"]
    hs2.font.size = Pt(14)
    hs2.font.bold = True

    hs3 = doc.styles["Heading 3"]
    hs3.font.size = Pt(12)
    hs3.font.bold = True

    # Code style
    code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.left_indent = Cm(1)


def add_code_block(doc, code_text):
    """Add a formatted code block."""
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph(line, style="CodeBlock")
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        p._p.get_or_add_pPr().append(shading)


def add_table_from_data(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "003366")
        run.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            if c_idx == 0:
                run.bold = True
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    doc.add_paragraph()
    return table


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p


def add_label_value(doc, label, value):
    """Add a label: value line."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(value)
    run.font.size = Pt(11)


# =========================================================================
# DOCUMENT 1: PROJECT DOCUMENTATION
# =========================================================================

def build_project_documentation():
    doc = Document()
    apply_styles(doc)

    # -- Cover Page --
    create_cover_page(
        doc,
        "Eco Rescue FPS",
        subtitle="Project Development Report",
        extra_lines=[
            ("Module:", "CM2121 — 3D Reconstructive Techniques"),
            ("Project:", "Eco Rescue FPS — SDG 12 Recycling Game"),
            ("Student:", "Matthew Jacob SD"),
            ("Student ID:", "[ID Placeholder]"),
            ("University:", "[University Name]"),
            ("Date:", "6 August 2026"),
            ("Extension Deadline:", "6 August 2026"),
        ],
    )

    # -- Table of Contents --
    p = doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # =====================================================================
    # 1. PROJECT OVERVIEW
    # =====================================================================
    doc.add_heading("1. Project Overview", level=1)

    doc.add_heading("1.1 Project Purpose", level=2)
    doc.add_paragraph(
        "Eco Rescue FPS is a first-person recycling simulation game built for the CM2121 "
        "resit assessment. The player collects recyclable items — plants, toys, and bottles "
        "— scattered across an open environment and disposes of them in the correct recycling "
        "bins within a 5-minute time limit, earning points and maintaining lives."
    )

    doc.add_heading("1.2 SDG Objective", level=2)
    doc.add_paragraph(
        "The project aligns with UN SDG 12: Responsible Consumption & Production. "
        "The core gameplay loop — collecting, sorting, and correctly recycling waste — "
        "reinforces the importance of proper waste management. Wrongly recycled items cost "
        "lives, teaching the consequence of incorrect disposal."
    )

    doc.add_heading("1.3 Final Gameplay", level=2)
    add_table_from_data(
        doc,
        ["Mechanic", "Description"],
        [
            ["Movement", "WASD movement, mouse look, sprint (Shift), jump (Space), storm wind push"],
            ["Interaction", "E to pick up, Q to drop, Left-click to throw, Right-click to aim"],
            ["Recycling", "Three bin types (Nature/Plastic/General Waste) with acceptance matrix"],
            ["Scoring", "Points per correct recycle, plant chain bonus, penalties for wrong bin"],
            ["Lives", "5 lives, lost on wrong-bin recycling"],
            ["Timer", "5-minute countdown, ends game when expired"],
            ["Weather", "3 states (Sunny/Rainy/Stormy); storm physically pushes player away from wrong bins"],
            ["HUD", "Score, lives, timer, collected items, announcements, score popups"],
            ["Menu", "Welcome → Instructions → Playing → End screen"],
            ["Pause", "ESC to pause, continue/settings/exit with save prompting"],
        ],
    )

    doc.add_heading("1.4 Current Architecture", level=2)
    add_code_block(doc, """Assets/
├── Audio/                 # 3 ambient tracks + 13 SFX clips
├── Environment/           # Terrain, water, lighting, weather prefabs
├── Input/                 # ActionsControl.inputactions (2 maps)
├── Models/                # 28 collectable prefabs + 3 bin prefabs
├── Scenes/
│   ├── Florance.unity     # Main game scene (build target)
│   └── ChainFragrance.unity  # UI prototyping scene
├── Scripts/
│   ├── Core/              # 4 scripts
│   ├── Player/            # 5 scripts
│   ├── Interaction/       # 2 scripts
│   ├── UI/                # 5 scripts
│   ├── Weather/           # 4 scripts + Data/ + Effects/
│   └── Editor/            # 1 script
├── Sounds/Optimized/      # Curated gameplay SFX + ambient clips (6 Aug remap)
├── UI/Sprites/            # GlassCard, Hail, Splats
├── Settings/              # URP pipeline assets
├── TextMesh Pro/          # Fonts, style sheets
├── Packages/              # manifest.json
└── ProjectSettings/       # Unity settings""")

    doc.add_paragraph(f"Script Count: 30 runtime scripts + 1 Editor script = 31 total.")

    # =====================================================================
    # 2. DEVELOPMENT TIMELINE
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("2. Development Timeline", level=1)

    doc.add_paragraph(
        "The following section provides a chronological overview of the project's "
        "evolution, explaining why each stage was introduced and what problems it solved."
    )

    doc.add_heading("2.1 Phase 1: Initial Prototype (3–15 July)", level=2)
    doc.add_paragraph(
        "The project began as a fork from the original resit-assessment prototype repository. "
        "This prototype established the core vision: a first-person recycling game with scanned "
        "photogrammetry assets."
    )
    doc.add_paragraph("Key achievements of the prototype phase:")
    add_bullet(doc, "Basic player movement (walk, sprint, jump)")
    add_bullet(doc, "Pickup, drop, and throw mechanics")
    add_bullet(doc, "9-zone grid layout for collectibles")
    add_bullet(doc, "Binary weather (sunny/rainy)")
    add_bullet(doc, "Legacy UI.Text-based HUD")

    doc.add_paragraph("Problems identified:")
    add_bullet(doc, "Procedural scene generation was inflexible and difficult to iterate on")
    add_bullet(doc, "Weather lacked dynamic transitions and felt binary")
    add_bullet(doc, "UI was basic and not event-driven; no game state flow")
    add_bullet(doc, "Game started immediately with no menu flow")

    doc.add_heading("2.2 Phase 2: Core Rebuild (15–22 July)", level=2)
    doc.add_paragraph(
        "The project was rebuilt from scratch in the CM2121 Resit Assessment - MJSD repository "
        "with a cleaner, modular architecture."
    )

    doc.add_heading("2.2.1 Project Init (15 July)", level=3)
    doc.add_paragraph(
        "Initial Unity project setup with Git repository initialisation and assessment "
        "requirements documentation. Established the project's directory structure."
    )

    doc.add_heading("2.2.2 Player Movement (16 July)", level=3)
    doc.add_paragraph(
        "CharacterController-based movement rewritten from scratch. Mouse look with "
        "sensitivity settings and invert-Y option. Jump physics with ground check. "
        "The original movement was tightly coupled; the new version uses modular components "
        "(PlayerMovement, PlayerLook) with event-driven APIs."
    )

    doc.add_heading("2.2.3 Folder Restructure and Asset Migration (19 July)", level=3)
    doc.add_paragraph(
        "Audio files migrated and organised, photogrammetry models imported, .gitignore "
        "configured. The original project had assets scattered across folders; these were "
        "restructured into clear categories (Audio, Models, Environment, Scripts)."
    )

    doc.add_heading("2.2.4 Interaction System (20 July)", level=3)
    doc.add_paragraph(
        "Pickup/drop/throw rewritten with smooth follow camera, layer and tag refactoring "
        "(Recyclable, NonRecyclable, RecyclingBin), single-item hold, and UI prompts with "
        "warnings. The original interaction was basic with no visual feedback; the new system "
        "uses smooth lerp follow, charge-based throw, and event-driven UI prompts."
    )

    doc.add_heading("2.3 Phase 3: Recycling and Scoring (21–22 July)", level=2)

    doc.add_heading("2.3.1 Recycling System (21 July)", level=3)
    doc.add_paragraph(
        "Three-bin types with acceptance matrix implemented. Score tracking with chain "
        "bonuses for consecutive correct recycles. Recycling events integrated into UI. "
        "The original had no proper recycling mechanic — items were simply picked up and "
        "dropped. The new system validates item type against bin type with point rewards "
        "and life penalties."
    )

    doc.add_heading("2.3.2 UI Integration (22 July)", level=3)
    doc.add_paragraph(
        "UIManager created with panel state machine (Welcome, Instructions, Playing, Ended). "
        "Input Action Map switching between Player and UI maps implemented. TextMesh Pro "
        "migration across all UI elements. The original had no menu flow; the new card-based "
        "panel system provides structured game flow with explicit state transitions."
    )

    doc.add_heading("2.4 Phase 4: Weather System Implementation (24 July)", level=2)
    doc.add_paragraph(
        "Proximity-based weather detection implemented via OverlapSphere. WeatherFeedbackSystem "
        "rewritten with three-state weather (Sunny, Rainy, Stormy) with storm intensity scaling. "
        "Per-state VFX pipeline created: clouds, rain, lightning, wind, and storm overlay. "
        "Ambient audio crossfade per weather state. The original weather was binary (on/off) "
        "triggered by pickup; the new system uses proximity detection near bins to dynamically "
        "scale weather intensity based on correct or incorrect recycling behaviour."
    )

    doc.add_heading("2.5 Phase 5: Audio and Polish (26 July)", level=2)
    doc.add_paragraph(
        "Surface-based footstep audio implemented with wet/dry detection. Water splash detection "
        "with drying timer (5 seconds after leaving water). Ambient audio crossfade between sunny, "
        "rainy and thunder tracks. The original had no footstep audio; the new system detects "
        "surface material (soil vs water) and plays appropriate sounds."
    )

    doc.add_heading("2.6 Phase 6: Pause Menu and UI Redesign (27–29 July)", level=2)

    doc.add_heading("2.6.1 ChainFragrance Scene (27 July)", level=3)
    doc.add_paragraph(
        "Dedicated UI prototyping scene created to avoid risking gameplay references in the "
        "main scene. Full PauseMenu hierarchy built: PausePanel, SettingsPanel, "
        "ConfirmationModal, and SaveProgressModal. Initial PauseMenuManager script created."
    )

    doc.add_heading("2.6.2 UI Migration (28 July)", level=3)
    doc.add_paragraph(
        "UIManager refactored with Pause action support. GameManager updated with PauseGame "
        "and ResumeGame methods. PauseMenuManager finalised with settings, exit flow, and "
        "username/volume configuration. AutoSetupPauseMenu Editor tool created for one-click "
        "scene setup. The original UIManager had no pause capability; the new system adds "
        "ESC-to-pause, settings panel, and exit confirmation with save prompting."
    )

    doc.add_heading("2.6.3 Final Polish (29 July)", level=3)
    doc.add_paragraph(
        "Removed duplicate broken UIManager from NewScripts/ folder. Editor tool updated to "
        "support any scene. Documentation updated with user testing logs. Script count "
        "finalised at 30 runtime plus one Editor script."
    )

    doc.add_heading("2.7 Phase 7: Documentation and Submission (30 July)", level=2)
    doc.add_paragraph(
        "Comprehensive project evolution report generated. Peer testing logs compiled with "
        "feedback from three testers. Professional Word documents prepared for submission. "
        "Extension deadline confirmed for 6 August 2026."
    )

    doc.add_heading("2.8 Phase 8: Resit Fix Pass and Final Hardening (3–6 August)", level=2)
    doc.add_paragraph(
        "Pre-submission audit pass fixed the broken core loop: colliders/layers on all "
        "collectibles and bins, GameManager scoring/lives/win-condition rewrite, and full UI "
        "rewiring via batch editor tools. A 21-test automated suite was added (EditMode + "
        "PlayMode) behind assembly definitions."
    )
    doc.add_paragraph(
        "Final hardening on 5–6 August made the environment self-contained (water + terrain "
        "materials moved off the demo folder; the 4.2 GB TerrainDemoScene_URP demo was removed), "
        "grounded the player on the flat terrain, added a storm wind-push consequence for wrong-bin "
        "recycling, remapped all gameplay SFX to the curated Optimized clips, corrected the bin "
        "acceptance types (all three prefabs were serialised as NatureRecycling), and cleaned up "
        "the HUD. Three batch validation passes completed with exit 0 and clean compilation."
    )

    # =====================================================================
    # 3. ARCHITECTURE EVOLUTION
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("3. Architecture Evolution", level=1)

    doc.add_heading("3.1 Folder Structure Evolution", level=2)
    doc.add_paragraph(
        "The project's folder structure underwent several reorganisations to improve "
        "maintainability and make project assets easier to locate."
    )

    doc.add_paragraph("Early structure (prototype):", style="Heading 3")
    add_code_block(doc, """Assets/Scripts/
├── GameManager.cs
├── PlayerController.cs
├── Interaction.cs
├── UIManager.cs
└── WeatherController.cs""")
    doc.add_paragraph(
        "Problems: All scripts flat in one folder. No separation of concerns. "
        "Difficult to maintain as the project grew."
    )

    doc.add_paragraph("Intermediate structure:", style="Heading 3")
    add_code_block(doc, """Assets/Scripts/
├── Core/              (GameManager, ScoreManager, AudioManager)
├── Player/            (Movement, Look, Interaction)
├── UI/                (UIManagers, HUDManager)
├── Weather/           (WeatherState, Effects)
└── NewScripts/        (UIManager.cs - prototype/broken)""")
    doc.add_paragraph(
        "Problems: NewScripts/ folder contained duplicate broken scripts. "
        "Poor naming consistency (UIManagers vs UIManager)."
    )

    doc.add_paragraph("Final structure:", style="Heading 3")
    add_code_block(doc, """Assets/Scripts/
├── Core/              # GameManager, ScoreManager, AudioManager, AutoSpawner
├── Player/            # Movement, Look, Interaction, FootstepAudio, WeatherEffect
├── Interaction/       # PickupItem, RecycleBinInteractable
├── UI/                # UIManager, HUDManager, PauseMenuManager, WeatherUI, InteractionPromptUI
├── Weather/           # State, FeedbackSystem, Effects, AnchorFollow
│   ├── Data/          # Parameters, SplashData
│   └── Effects/       # Wind, Cloud, Rain, Sunny, Lightning, Splash
└── Editor/            # AutoSetupPauseMenu""")
    doc.add_paragraph(
        "Benefits: Clear separation by system. Editor scripts isolated from runtime code. "
        "No duplicate managers. Intuitive navigation for future development."
    )

    doc.add_heading("3.2 UI Organisation Evolution", level=2)
    doc.add_paragraph(
        "The UI evolved from an immediate-start flow with legacy text to a structured "
        "card-based panel system with TextMesh Pro and pause overlay support."
    )
    add_table_from_data(
        doc,
        ["Phase", "Structure", "Technologies"],
        [
            ["Prototype", "No menu, immediate game start", "Legacy UI.Text"],
            ["Intermediate", "4 panels (Welcome, Instructions, Playing, Ended)", "TextMesh Pro migration began"],
            ["Final", "4 main panels + 4 pause sub-panels", "TextMesh Pro, event-driven, pause overlay"],
        ],
    )

    doc.add_heading("3.3 Scene Structure Evolution", level=2)
    add_table_from_data(
        doc,
        ["Phase", "Scene Configuration", "Purpose"],
        [
            ["Prototype", "Single scene, procedural generation", "Core gameplay testing"],
            ["Intermediate", "Florance.unity (main), ChainFragrance.unity (add)", "Gameplay + UI prototyping separation"],
            ["Final", "Florance.unity (build target), ChainFragrance.unity (reference)", "Production + UI sandbox"],
        ],
    )

    doc.add_heading("3.4 Manager Restructuring", level=2)
    add_table_from_data(
        doc,
        ["Original", "Problem", "Final"],
        [
            ["GameManager (monolithic)", "Too many responsibilities", "GameManager + ScoreManager + AutoSpawner"],
            ["UIManager (flat)", "No pause support", "UIManager + PauseMenuManager"],
            ["No editor tools", "Manual setup error-prone", "AutoSetupPauseMenu Editor tool"],
        ],
    )

    # =====================================================================
    # 4. SCRIPT EVOLUTION
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("4. Script Evolution", level=1)
    doc.add_paragraph(
        "For every major system, this section explains the original implementation, "
        "problems encountered, changes made, and the final implementation."
    )

    doc.add_heading("4.1 GameManager", level=2)
    doc.add_paragraph(
        "Original (prototype): Monolithic class handling game state, scoring, lives, timer, "
        "and restart in a single file. No event system — directly manipulated UI elements."
    )
    doc.add_paragraph("Problems:")
    add_bullet(doc, "Tightly coupled with UI scripts")
    add_bullet(doc, "No separation of scoring logic")
    add_bullet(doc, "No pause/resume support")
    add_bullet(doc, "RestartGame() did not reset Time.timeScale")
    doc.add_paragraph("Changes made:")
    add_bullet(doc, "Score logic extracted to ScoreManager with event-driven updates")
    add_bullet(doc, "Added PauseGame() and ResumeGame() methods to stop/resume timer")
    add_bullet(doc, "RestartGame() now resets Time.timeScale to 1f")
    add_bullet(doc, "All UI communication via C# events (OnGameStarted, OnGameOver, etc.)")
    add_bullet(doc, "3 Aug rewrite: signed scoring via ReportRecycled, 12/8/4 per-category win objectives, lives only lost on negative scores")
    doc.add_paragraph("Final: 324 lines. Handles game state, lives, timer, recycling reporting, "
                       "win/loss checks. Event-driven and decoupled from UI.")

    doc.add_heading("4.2 UIManager", level=2)
    doc.add_paragraph(
        "Original: UIManagers.cs (plural) managed panel states but had no pause support. "
        "NewScripts/UIManager.cs was a broken duplicate with an infinite while-loop and empty stubs."
    )
    doc.add_paragraph("Problems:")
    add_bullet(doc, "Class name UIManagers inconsistent with project conventions")
    add_bullet(doc, "Duplicate class caused namespace conflicts")
    add_bullet(doc, "No pause action handling")
    add_bullet(doc, "Input map switching did not accommodate pause overlay")
    doc.add_paragraph("Changes made:")
    add_bullet(doc, "Renamed to UIManager.cs (singular)")
    add_bullet(doc, "Added pauseMenuPanel GameObject reference")
    add_bullet(doc, "Added pauseAction from UI action map, selectively enabled during gameplay")
    add_bullet(doc, "Added ShowPauseMenu() and ReturnFromPause() public methods")
    add_bullet(doc, "Input map switching now caches playerMap/uiMap for reuse")
    doc.add_paragraph(
        "Final: UIManager owns the main game flow panel state machine. Delegates pause "
        "menu logic to PauseMenuManager."
    )

    doc.add_heading("4.3 PauseMenuManager", level=2)
    doc.add_paragraph(
        "Created on 27 July 2026 as a dedicated pause menu controller, separate from the main "
        "UI flow. 211 lines."
    )
    doc.add_paragraph("Features:")
    add_bullet(doc, "ESC key toggles pause (via UIManager coordination)")
    add_bullet(doc, "Time.timeScale = 0 freezes game physics and timers")
    add_bullet(doc, "Continue button resumes gameplay")
    add_bullet(doc, "Settings panel with username (PlayerPrefs) and volume slider (AudioListener.volume)")
    add_bullet(doc, "Two-step exit flow: Confirm → Save prompt → Quit (reloads scene to main menu)")
    add_bullet(doc, "Restart button reloads scene")
    doc.add_paragraph(
        "Architecture note: PauseMenuManager calls UIManager methods for input map "
        "switching and cursor management, avoiding duplication of input handling logic."
    )

    doc.add_heading("4.4 Player Controller (PlayerMovement + PlayerLook)", level=2)
    doc.add_paragraph(
        "Original (prototype): Single PlayerController.cs with movement, look, sprint, jump, "
        "crouch, slide, dash, and trip mechanics — over 400 lines."
    )
    doc.add_paragraph("Problems:")
    add_bullet(doc, "400+ line monolith difficult to maintain")
    add_bullet(doc, "Too many movement modes (slide, dash, trip) unused in final game")
    add_bullet(doc, "Speed modifiers hardcoded")
    add_bullet(doc, "No public API for weather effects")
    doc.add_paragraph("Changes made:")
    add_bullet(doc, "Split into PlayerMovement.cs (190 lines) and PlayerLook.cs (83 lines)")
    add_bullet(doc, "Removed unused movement modes")
    add_bullet(doc, "Added SetSpeedModifier(float) public API for weather system")
    add_bullet(doc, "Added SetWindPush(Vector3) — storm wrong-bin push applied via CharacterController")
    doc.add_paragraph(
        "Final: Clean, focused components. PlayerMovement handles WASD, sprint, jump, "
        "and wind push. PlayerLook handles mouse with sensitivity and invert-Y."
    )

    doc.add_heading("4.5 Player Interaction", level=2)
    doc.add_paragraph(
        "Original (prototype): Basic raycast pickup with no visual feedback. Drop and throw "
        "were toggle-based."
    )
    doc.add_paragraph("Changes made:")
    add_bullet(doc, "Smooth lerp follow for held items (instead of snap-to)")
    add_bullet(doc, "Charge-based throw (hold left mouse to charge)")
    add_bullet(doc, "Layer filtering (only Interactable layer)")
    add_bullet(doc, "Event-driven UI prompts via InteractionPromptUI")
    add_bullet(doc, "Distance and angle validation")
    doc.add_paragraph(
        "Final: 211 lines. Robust interaction with visual feedback. Items smoothly follow "
        "camera. Throw power scales with charge time."
    )

    doc.add_heading("4.6 Recycling System (PickupItem + RecycleBinInteractable)", level=2)
    doc.add_paragraph(
        "Original (prototype): Nine bins with no type checking. Any item could go in any bin."
    )
    doc.add_paragraph("Problems:")
    add_bullet(doc, "No gameplay consequence for incorrect disposal")
    add_bullet(doc, "9-bin grid confusing for players")
    add_bullet(doc, "Bin prefab serialised types all set to NatureRecycling (fixed 6 Aug)")
    doc.add_paragraph("Changes made:")
    add_bullet(doc, "PickupItem.cs: Item component with ItemType enum (Plant/Toy/Bottle)")
    add_bullet(doc, "RecycleBinInteractable.cs: Bin component with BinType enum (Nature/Plastic/General)")
    add_bullet(doc, "3x3 acceptance matrix with point rewards and life penalties")
    add_bullet(doc, "Storm wind-push penalty: wrong-bin recycling pushes the player away")

    doc.add_paragraph("Acceptance Matrix (signed score per recycle; positive = correct):")
    add_table_from_data(
        doc,
        ["Item / Bin", "Nature Recycling", "Plastic Recycling", "General Waste"],
        [
            ["Plant", "+20", "-45", "-20"],
            ["Bottle", "-15", "+20", "+15"],
            ["Toy", "-25", "-15", "+25"],
        ],
    )
    doc.add_paragraph(
        "Correct recycles add the (positive) score; wrong recycles subtract it and cost a life. "
        "Plant chain bonus: 2 consecutive correct plants = +40 bonus points."
    )

    doc.add_heading("4.7 Weather System", level=2)
    doc.add_paragraph(
        "Original (prototype): Binary sunny/rainy toggle via string. No particles, "
        "no transitions, no visual effects."
    )
    doc.add_paragraph("Problems:")
    add_bullet(doc, "Weather did not feel dynamic")
    add_bullet(doc, "No visual feedback (just text)")
    add_bullet(doc, "No gameplay impact")
    doc.add_paragraph("Final implementation components:")
    add_table_from_data(
        doc,
        ["Component", "Lines", "Purpose"],
        [
            ["WeatherState", "37", "3-state enum (Sunny/Rainy/Stormy) with C# events"],
            ["WeatherFeedbackSystem", "181", "Proximity detection via OverlapSphere, storm intensity scaling"],
            ["WeatherEffects", "174", "VFX orchestration across 6 effect types"],
            ["CloudEffect", "34", "Cloud particle color and emission rate"],
            ["RainEffect", "30", "Rain emission rate (200-800 particles)"],
            ["LightingEffect + LightningFlash", "79", "Random lightning flashes with VFX"],
            ["WindEffect", "67", "WindZone + wind particles, per-state speed"],
            ["SunnyEffect", "31", "Sun intensity + god rays particles"],
            ["WeatherMovementEffect", "62", "Speed modifiers per state"],
            ["WeatherAnchorFollow", "19", "VFX anchor follows player"],
        ],
    )
    doc.add_paragraph(
        "Weather feedback is triggered by proximity to bins and player recycling behaviour: "
        "correct recycles calm the weather, wrong recycles escalate to a storm that physically "
        "pushes the player away from the bin (max 3 m/s, ramped)."
    )

    doc.add_heading("4.8 Audio System", level=2)
    doc.add_paragraph(
        "Original (prototype): Single AudioSource with crossfade between sunny/rainy tracks. "
        "No SFX, no footsteps."
    )
    doc.add_paragraph("Changes made:")
    add_bullet(doc, "AudioManager.cs (159 lines): Dual-source ambient crossfade between sunny, rainy, thunder")
    add_bullet(doc, "SFX playback method for gameplay sounds")
    add_bullet(doc, "PlayerFootstepAudio.cs (202 lines): Surface-based detection with 5s drying timer")
    add_bullet(doc, "Splash ParticleSystem spawning on water surfaces")
    add_bullet(doc, "6 Aug remap: all gameplay SFX (footsteps + bin success/error) reference the curated Sounds/Optimized clips")
    doc.add_paragraph("Total audio library: 19 clips (3 ambient + 13 SFX + 3 weather-specific)")

    doc.add_heading("4.9 HUD", level=2)
    doc.add_paragraph(
        "Original (prototype): Legacy UnityEngine.UI.Text elements with no event-driven "
        "updates. Texts updated by polling GameManager each frame."
    )
    doc.add_paragraph("Changes made:")
    add_bullet(doc, "TextMesh Pro migration across all HUD elements")
    add_bullet(doc, "Event-driven updates via GameManager/ScoreManager events")
    add_bullet(doc, "Announcement system (timed text with 3s duration)")
    add_bullet(doc, "Score popup system (3 popup GameObjects by item type)")
    add_bullet(doc, "Programmatic fallback text creation for missing references")
    add_bullet(doc, "Timer colour coding: white (>60s), yellow (30-60s), red (<30s)")
    add_bullet(doc, "6 Aug: HUD writes bare numeric values matching the static scene header labels")
    doc.add_paragraph("Final: 273 lines. Modern, event-driven HUD with fallback creation.")

    doc.add_heading("4.10 Scene Management", level=2)
    doc.add_paragraph(
        "Original: Single scene with no build settings configured. No restart capability."
    )
    doc.add_paragraph("Changes made:")
    add_bullet(doc, "Build settings: Florance.unity at index 0")
    add_bullet(doc, "GameManager.RestartGame() uses SceneManager.LoadScene() with scene reload")
    add_bullet(doc, "Managers use DontDestroyOnLoad to persist across restarts")
    add_bullet(doc, "AutoSpawner.cs safety-net script instantiates missing managers if absent")

    # =====================================================================
    # 5. LATEST PROJECT STATE
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("5. Latest Project State", level=1)

    doc.add_heading("5.1 Runtime Scripts (30 total)", level=2)
    add_table_from_data(
        doc,
        ["Category", "Script", "Lines", "Purpose"],
        [
            ["Core", "GameManager", "324", "Game state, lives, timer, recycling"],
            ["Core", "ScoreManager", "103", "Score tracking, high score"],
            ["Core", "AudioManager", "159", "Dual-source ambient, SFX"],
            ["Core", "AutoSpawner", "47", "Safety-net manager spawner"],
            ["Player", "PlayerMovement", "190", "WASD, sprint, jump, speed + wind push API"],
            ["Player", "PlayerLook", "83", "Mouse look, sensitivity"],
            ["Player", "PlayerInteraction", "261", "Pickup/drop/throw"],
            ["Player", "PlayerFootstepAudio", "202", "Surface-based footsteps"],
            ["Player", "WeatherMovementEffect", "92", "Weather speed modifiers"],
            ["Interaction", "PickupItem", "125", "Item type, physics"],
            ["Interaction", "RecycleBinInteractable", "132", "Bin type, recycling trigger"],
            ["UI", "UIManager", "285", "Panel state machine, input maps"],
            ["UI", "HUDManager", "273", "Stats, announcements, popups"],
            ["UI", "PauseMenuManager", "236", "Pause, settings, exit flow"],
            ["UI", "WeatherUI", "146", "Weather display"],
            ["UI", "InteractionPromptUI", "116", "Pickup/drop prompts"],
            ["Weather", "WeatherState", "75", "Weather enum + events"],
            ["Weather", "WeatherFeedbackSystem", "260", "Proximity detection + storm wind push"],
            ["Weather", "WeatherEffects", "220", "VFX orchestration"],
            ["Weather", "WeatherAnchorFollow", "32", "VFX follow player"],
            ["Weather/Data", "WeatherEffectParameters", "45", "VFX config"],
            ["Weather/Data", "SplashData", "25", "Splash config"],
            ["Weather/Effects", "WindEffect", "93", "WindZone + particles"],
            ["Weather/Effects", "CloudEffect", "53", "Cloud particles"],
            ["Weather/Effects", "RainEffect", "53", "Rain particles"],
            ["Weather/Effects", "SunnyEffect", "54", "Sun + god rays"],
            ["Weather/Effects", "LightingEffect", "40", "Lightning wrapper"],
            ["Weather/Effects", "LightingFlash", "83", "Random lightning"],
            ["Weather/Effects", "SplashEffect", "46", "Splash particles"],
            ["Weather/Effects", "SplashSpawner", "27", "Spawn splash prefabs"],
        ],
    )

    doc.add_heading("5.2 UI Hierarchy", level=2)
    add_code_block(doc, """Canvas (Screen Space - Overlay, 1280x720)
├── PanelUI/
│   ├── WelcomeScreen          # "Press Enter to continue"
│   ├── InstructionScreen      # Controls guide
│   ├── HUDScreen              # Lives, score, timer, items, popups
│   └── CreditsScreen          # End screen with results
└── PauseMenu/                 # Toggled by ESC during gameplay
    ├── PausePanel             # Continue / Settings / Exit
    ├── SettingsPanel          # Username input + Volume slider
    ├── ConfirmationModal      # "Are you sure you want to exit?"
    └── SaveProgressModal      # "Save before quitting?" """)

    doc.add_heading("5.3 Input System", level=2)
    doc.add_paragraph("Action Maps:")
    add_table_from_data(
        doc,
        ["Map", "Actions", "Bindings"],
        [
            ["Player", "Movement, Look, Jump, Sprint, Interact, Drop, Throw, Aim", "WASD, Mouse, Space, Shift, E, Q, LMB, RMB"],
            ["UI", "Navigate, Submit, Cancel, Pause, Restart, Continue, Scroll", "WASD/Arrows, Enter/Space, Escape, R, Enter, Scroll"],
        ],
    )

    doc.add_paragraph("Map Switching Logic:", style="Heading 3")
    add_table_from_data(
        doc,
        ["State", "Player Map", "UI Map", "Pause Action"],
        [
            ["Welcome", "Disabled", "Enabled", "Disabled"],
            ["Instructions", "Disabled", "Enabled", "Disabled"],
            ["Playing", "Enabled", "Disabled", "Enabled (selective)"],
            ["Playing → Paused", "Disabled", "Enabled", "Disabled"],
            ["Ended", "Disabled", "Enabled", "Disabled"],
        ],
    )

    doc.add_heading("5.4 Gameplay Flow", level=2)
    add_code_block(doc, """Scene Load → Welcome Screen
                  ↓ (Enter key)
            Instruction Screen
                  ↓ (Enter key)
            Game Starts (5:00 timer)
                  ↓
        ┌─── Playing (HUD visible) ───┐
        │         ↕ (ESC)              │
        │     Pause Menu               │
        │   ├── Continue ────────┐     │
        │   ├── Settings          │     │
        │   └── Exit → Confirm    │     │
        │            → Save? → Quit     │
        │                              │
        ├── Timer expires              │
        ├── All plants recycled (Won)  │
        └── Lives depleted (Lost)      │
                  ↓
            End Screen (Credits)
                  ↓ (Restart)
            Scene reload → Welcome""")

    doc.add_heading("5.5 Known Limitations", level=2)
    add_table_from_data(
        doc,
        ["Issue", "Impact", "Workaround"],
        [
            ["8 toy instances + WaterPlane stacked off-world", "Some toys unreachable until repositioned", "Reposition in Unity Editor before final play"],
            ["Legacy RawAudio originals untracked", "Repo clone lacks raw sources", "Optimized clips are the referenced set; raw kept for provenance"],
            ["Audio volume inconsistent", "Some clips louder than others", "Manual volume adjustment in AudioSource clips"],
            ["Collectables spread wide", "Hard to find all items", "Increase game timer or reposition items"],
            ["No post-processing volume", "Scene lacks aesthetic tuning", "Add URP Volume with bloom, tone mapping"],
        ],
    )

    doc.add_heading("5.6 Auto Setup Workflow", level=2)
    doc.add_paragraph(
        "The AutoSetupPauseMenu Editor tool (Tools → Setup UI in Current Scene) automates:\n"
        "1. Canvas structure creation (PanelUI hierarchy if missing)\n"
        "2. UIManager creation with panel references wired\n"
        "3. PauseMenuManager with pause panels, images, and VerticalLayoutGroups\n"
        "4. Button OnClick rewiring to correct PauseMenuManager methods\n"
        "5. Auto-detection of ActionsControl.inputactions asset\n"
        "6. HUDManager TMP text reference wiring"
    )
    doc.add_paragraph("Post-tool manual steps:")
    add_bullet(doc, "Add VolumeSlider to SettingsPanel")
    add_bullet(doc, "Verify HUDManager references")

    # =====================================================================
    # 6. COMPLETE CHANGE LOG
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("6. Complete Change Log", level=1)
    doc.add_paragraph(
        "The following change log summarises the major modifications made throughout "
        "development, including what changed, why, which files were affected, benefits "
        "introduced, and any trade-offs."
    )

    changes = [
        ("UIManager Refactoring",
         "Class renamed from UIManagers to UIManager. Added pause support with "
         "ShowPauseMenu/ReturnFromPause. Cached input maps. Selective Pause action enabling.",
         "Name inconsistency; needed pause support without breaking existing flow.",
         "Assets/Scripts/UI/UIManager.cs",
         "Clean API for pause integration; no duplicate map lookups.",
         "Existing scenes need pauseMenuPanel wired."),
        ("PauseMenuManager Introduction",
         "New script PauseMenuManager.cs created with pause menu, settings, and exit flow.",
         "Separate pause logic from main UI flow for focused components.",
         "Assets/Scripts/UI/PauseMenuManager.cs (new)",
         "Clean separation of concerns; settings isolated; full undo support.",
         "Additional script to maintain."),
        ("AutoSetupPauseMenu Tool",
         "Editor tool that automates scene UI setup with one click.",
         "Manual Inspector wiring is error-prone and time-consuming.",
         "Assets/Scripts/Editor/AutoSetupPauseMenu.cs (new)",
         "One-click setup; reduces human error; handles both scenes; undo support.",
         "Must be run once per scene after creation."),
        ("Removal of Duplicate UIManager",
         "Deleted Assets/Scripts/NewScripts/UIManager.cs.",
         "Duplicate class caused namespace conflict; script was broken with infinite loop.",
         "Assets/Scripts/NewScripts/UIManager.cs (deleted)",
         "Resolved compilation ambiguity; removed dead code; cleaner project structure.",
         "None."),
        ("GameManager Updates",
         "Added PauseGame(), ResumeGame(), and Time.timeScale reset in RestartGame().",
         "Pause menu needed to freeze timer without ending the game.",
         "Assets/Scripts/Core/GameManager.cs",
         "Pause functions without game-ending side effects; timer resumes correctly.",
         "None."),
        ("GameManager Timer Fix",
         "Removed incorrect OnPauseTime assignment that reset timer to 300.",
         "Pause was resetting timer to full duration instead of stopping countdown.",
         "Assets/Scripts/Core/GameManager.cs",
         "Timer accurately tracks remaining time across pause/resume cycles.",
         "None."),
        ("Documentation Updates",
         "Created docs/USER_TESTING.md and docs/PROJECT_EVOLUTION.md. Updated "
         "ASSESSMENT_REQUIREMENTS.md and TIMELINE.md.",
         "Assessment requires evidence of testing and full project documentation.",
         "Multiple files across docs/",
         "Full testing documentation for submission; peer feedback catalogued.",
         "Time investment in writing."),
        ("Scene Migration",
         "Pause menu hierarchy and scripts moved from ChainFragrance to Florance.",
         "ChainFragrance was a sandbox; Florance is the build target.",
         "Both .unity scene files",
         "Main scene has full pause menu; prototyping scene preserved for reference.",
         "Inspector references needed rewiring after migration."),
        ("UI Redesign",
         "Card-based panel system replaced immediate-start flow.",
         "Players need clear onboarding before gameplay.",
         "Assets/Scripts/UI/UIManager.cs, scene files",
         "Structured onboarding; clear game states; proper cursor management.",
         "Additional scene setup required."),
        ("HUD Redesign",
         "Legacy Text replaced with TextMesh Pro. Event-driven updates. Announcement "
         "and popup systems added.",
         "Legacy Text lacked quality; polling was wasteful.",
         "Assets/Scripts/UI/HUDManager.cs",
         "Crisp text; performance-efficient events; rich feedback; fallback creation.",
         "Some references must be manually wired in Inspector."),
        ("Player Grounding Fix",
         "CharacterController center moved from (0,1,0) to (0,0,0); groundMask widened from "
         "80 to 81 so the ground check hits the terrain layers.",
         "Player floated above the flat terrain and could not stand on it.",
         "Assets/Scenes/Florance.unity",
         "Feet flush with ground; jump and movement work on the play area.",
         "None."),
        ("Storm Wind Push",
         "WeatherFeedbackSystem now ramps a MoveTowards push (max 3 m/s, 2.5 m/s^2) away from "
         "a bin after a wrong recycle, applied via PlayerMovement.SetWindPush.",
         "Wrong-bin recycling had no physical consequence; storm was purely visual.",
         "Assets/Scripts/Weather/WeatherFeedbackSystem.cs, Assets/Scripts/Player/PlayerMovement.cs",
         "Clear gameplay consequence for wrong disposal; makes storm meaningful.",
         "Slightly more complex player controller API."),
        ("Gameplay Audio Remap",
         "Player footsteps (dry/run/wet) and all bin success/error SFX remapped to the curated "
         "Sounds/Optimized clips (SFX_DryWalk, SFX_Running, SFX_WetWalk, SFX_Correct, SFX_Buzzer).",
         "Gameplay SFX referenced legacy RawAudio originals.",
         "Assets/Scenes/Florance.unity, Assets/Models/Prefabs/Bins/*.prefab",
         "Consistent curated SFX; zero RawAudio references remain.",
         "Optimized clips must be delivered with the repo."),
        ("Bin Type Correction",
         "All three bin prefabs were serialised as binType 0 (NatureRecycling); corrected to "
         "GeneralWaste (2) and PlasticRecycling (1).",
         "Plastic bottles/glass could never score; only plants were accepted anywhere.",
         "Assets/Models/Prefabs/Bins/*.prefab",
         "Acceptance matrix works as designed; verified against RecycleBinMatrixTests.",
         "None."),
        ("HUD Numeric Cleanup",
         "HUD score/lives write bare values matching the static scene header labels.",
         "HUD duplicated its own labels next to scene headers.",
         "Assets/Scripts/UI/HUDManager.cs",
         "Clean single label per stat.",
         "None."),
        ("Demo Folder Removal",
         "Removed the 4.2 GB Assets/TerrainDemoScene_URP/ folder (git rm -r).",
         "Game must build from tracked, self-contained assets only.",
         "Assets/TerrainDemoScene_URP/** (deleted)",
         "Zero demo-folder references remain; repository no longer depends on it.",
         "None."),
    ]

    for title, what, why, files, benefits, tradeoffs in changes:
        doc.add_heading(f"6.{changes.index((title, what, why, files, benefits, tradeoffs)) + 1} {title}", level=2)
        add_label_value(doc, "What", what)
        add_label_value(doc, "Why", why)
        add_label_value(doc, "Files affected", files)
        add_label_value(doc, "Benefits", benefits)
        add_label_value(doc, "Trade-offs", tradeoffs)

    # =====================================================================
    # 7. LESSONS LEARNED
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("7. Lessons Learned", level=1)

    lessons = [
        ("Unity Architecture",
         [
             "Separation of concerns matters: Splitting monolithic GameManager into GameManager + ScoreManager made debugging easier and allowed event-driven UI updates.",
             "Singleton pattern is useful but must be consistent: DontDestroyOnLoad singletons persist across scene reloads, but must reset state in StartGame() rather than Awake().",
             "Events over direct references: C# events (event Action) decouple UI from gameplay logic. HUDManager subscribes to GameManager events instead of being polled.",
         ]),
        ("Scene Management",
         [
             "Prototyping scenes are valuable: ChainFragrance.unity allowed isolated UI testing without risking gameplay references.",
             "Scene reload is simpler than additive loading: SceneManager.LoadScene resets all state cleanly, combined with DontDestroyOnLoad singletons for persistent data.",
             "Scene roots should be organised: Grouping Managers, Environment, UI under parent transforms keeps the hierarchy navigable.",
         ]),
        ("Inspector Configuration",
         [
             "Serialized fields are fragile: Missing references cause silent fallbacks. Always verify references in the Inspector.",
             "Editor tools reduce setup errors: AutoSetupPauseMenu eliminates 15+ manual Inspector drags and prevents mis-wired buttons.",
             "Null checks on serialized fields: if (field != null) before every use prevents NullReferenceExceptions in production builds.",
         ]),
        ("Prefab Organisation",
         [
             "Prefabs for collectables are essential: 28 collectable prefabs make scene population straightforward.",
             "No mixing of prefab types in a single folder: Separate folders per type (Bottle, Nature, Toys) simplifies asset management.",
         ]),
        ("UI Workflow",
         [
             "Panel state machines work well: Enum-based state switching is simple and readable.",
             "Pause overlay is tricky: It must overlay on the HUD without being a new state. Using ShowPauseMenu() (not ShowPanel(Paused)) keeps the HUD active underneath.",
             "Selective input map enabling: Enabling just the Pause action from a disabled UI map is cleaner than modifying action map bindings at runtime.",
         ]),
        ("Input System",
         [
             "Action Map switching is reliable: The new Input System's Enable()/Disable() per map prevents conflicting inputs between Player and UI maps.",
             "Selective action enabling is powerful: pauseAction.Enable() works even when the parent UI map is disabled.",
             "Cached maps improve performance: FindActionMap() called once in Awake() rather than every state change.",
         ]),
        ("Script Organisation",
         [
             "Folder-by-system structure scales well: Core, Player, Interaction, UI, Weather — each system has clear ownership.",
             "Editor scripts belong in Editor/ folder: Prevents Editor-only code from appearing in builds.",
             "Class naming consistency is important: Avoid duplicates across folders (UIManager vs UIManagers).",
         ]),
        ("Testing and Debugging",
         [
             "Peer testing reveals blind spots: Rose identified unclear objectives, McJames noted object placement issues, Gideon found jump and weather bugs.",
             "Systematic testing is better than ad-hoc: A checklist (movement, interaction, UI, audio, weather) ensures coverage.",
             "Unity Console is first line of defence: Missing references, null exceptions, and compilation errors all surface there.",
             "Debug.Log for event flow: Logging when events fire helps trace event chain issues.",
         ]),
        ("Documentation",
         [
             "Document as you go, not at the end: Writing docs at the end misses intermediate architecture decisions.",
             "Change logs capture evolution: Knowing why something changed is more valuable than knowing what changed.",
             "Timeline with dates grounds the project: Showing development duration demonstrates scope and effort.",
         ]),
    ]

    for title, points in lessons:
        doc.add_heading(f"7.{lessons.index((title, points)) + 1} {title}", level=2)
        for point in points:
            add_bullet(doc, point)

    # =====================================================================
    # 8. FUTURE IMPROVEMENTS
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("8. Future Improvements and Recommendations", level=1)
    doc.add_paragraph(
        "The following improvements are recommended for future development iterations:"
    )

    improvements = [
        "Resolve weather system triggering and transitions",
        "Improve environmental readability and player guidance",
        "Balance audio levels across all clips and weather states",
        "Add additional environmental feedback for correct/incorrect actions",
        "Improve objective guidance for players during gameplay",
        "Refine UI animations and transition effects between panels",
        "Increase environmental variety with additional biomes or areas",
        "Improve recycling feedback with particle effects and sounds",
        "Optimise collectable object placement density",
        "Continue refactoring scripts into a modular architecture",
        "Perform broader playtesting with additional users",
        "Optimise overall project performance and draw calls",
        "Improve visual polish and gameplay presentation",
        "Add post-processing volume for enhanced visual quality",
        "Implement a save/load system for full game state persistence",
    ]
    for imp in improvements:
        add_bullet(doc, imp)

    # -- Save --
    path = os.path.join(OUTPUT_DIR, "CM2121_Project_Documentation.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# =========================================================================
# DOCUMENT 2: TEST LOG
# =========================================================================

def build_test_log():
    doc = Document()
    apply_styles(doc)

    # -- Cover Page --
    create_cover_page(
        doc,
        "Eco Rescue FPS",
        subtitle="User Testing Report",
        extra_lines=[
            ("Module:", "CM2121 — 3D Reconstructive Techniques"),
            ("Project:", "Eco Rescue FPS — SDG 12 Recycling Game"),
            ("Student:", "Matthew Jacob SD"),
            ("Student ID:", "[ID Placeholder]"),
            ("Date:", "6 August 2026"),
        ],
    )

    # -- Table of Contents --
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # =====================================================================
    # 1. INTRODUCTION
    # =====================================================================
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document presents the testing outcomes for the Eco Rescue FPS project, "
        "a first-person recycling simulation game developed for the CM2121 module. "
        "Testing was conducted to evaluate gameplay functionality, user experience, "
        "and system reliability prior to final submission."
    )
    doc.add_paragraph(
        "The testing process involved both developer-led systematic testing and peer "
        "testing with three independent testers. Each tester evaluated the game across "
        "multiple criteria including movement, interaction, audio, weather, UI, and "
        "overall gameplay experience."
    )

    # =====================================================================
    # 2. TESTING METHODOLOGY
    # =====================================================================
    doc.add_heading("2. Testing Methodology", level=1)

    doc.add_heading("2.1 Developer Testing", level=2)
    doc.add_paragraph(
        "Developer testing was conducted throughout the development lifecycle using "
        "a systematic checklist. Each system (movement, interaction, recycling, weather, "
        "audio, UI) was tested individually in the Unity Editor play mode. The Unity "
        "Console was monitored for errors, warnings, and null reference exceptions "
        "during all testing sessions."
    )
    doc.add_paragraph(
        "A 21-test automated suite (EditMode + PlayMode) verifies scoring, lives, win/loss "
        "conditions, the bin acceptance matrix, and pause behaviour. Final hardening was "
        "validated with three batch compile passes (exit 0, clean compilation) covering the "
        "wind-push, audio remap, and HUD/bin-type changes."
    )

    doc.add_heading("2.2 Peer Testing", level=2)
    doc.add_paragraph(
        "Three independent peer testers evaluated the game in a controlled environment. "
        "Each tester played the game without prior instruction to assess whether the "
        "gameplay was intuitive. Testers were asked to provide feedback on:"
    )
    add_bullet(doc, "Movement and controls")
    add_bullet(doc, "Visual quality and environment design")
    add_bullet(doc, "Audio quality and appropriateness")
    add_bullet(doc, "Gameplay clarity and objective understanding")
    add_bullet(doc, "Overall enjoyment and engagement")
    add_bullet(doc, "Bugs or issues encountered")

    doc.add_heading("2.3 Functional Testing", level=2)
    doc.add_paragraph(
        "Functional testing verified that each core mechanic operated as specified:"
    )
    add_bullet(doc, "Player movement (walk, sprint, jump) in all weather conditions")
    add_bullet(doc, "Object interaction (pickup, drop, throw) with correct physics behaviour")
    add_bullet(doc, "Recycling mechanics with correct acceptance matrix and scoring")
    add_bullet(doc, "Timer countdown, lives tracking, and game over/win conditions")
    add_bullet(doc, "Weather state transitions and visual effects")
    add_bullet(doc, "UI panel navigation and button functionality")
    add_bullet(doc, "Pause menu, settings, and exit confirmation flow")
    add_bullet(doc, "Audio playback and crossfade behaviour")

    doc.add_heading("2.4 UI Testing", level=2)
    doc.add_paragraph(
        "UI testing verified that all interface elements were responsive, correctly "
        "positioned, and displayed accurate information. This included:"
    )
    add_bullet(doc, "Welcome, instruction, and end screen panels")
    add_bullet(doc, "HUD elements (score, lives, timer, item counts)")
    add_bullet(doc, "Pause menu and all sub-panels")
    add_bullet(doc, "Settings panel (username input, volume slider)")
    add_bullet(doc, "Confirmation and save progress dialogs")
    add_bullet(doc, "Interaction prompts and feedback notifications")

    doc.add_heading("2.5 Gameplay Testing", level=2)
    doc.add_paragraph(
        "Gameplay testing assessed the overall player experience including game flow, "
        "difficulty balance, objective clarity, and engagement. This involved playing "
        "through complete game sessions from welcome screen to end screen."
    )

    # =====================================================================
    # 3. PEER TESTING
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("3. Peer Testing Results", level=1)

    # -- Rose --
    doc.add_heading("3.1 Tester: Rose", level=2)

    add_table_from_data(
        doc,
        ["Category", "Rating / Feedback"],
        [
            ["Overall Rating", "6/10"],
            ["Movement", "Works correctly"],
            ["Visual Style", "Cute and appealing environment"],
            ["Functionality", "Project does not feel fully functional"],
            ["Bugs", "Movement bugs encountered"],
            ["Clarity", "Objective unclear despite instructions being present"],
            ["Audio", "Inconsistent and unbalanced"],
            ["Environment", "Lacks guidance in guiding the player"],
        ],
    )
    doc.add_paragraph(
        "Summary: Rose felt the project showed good effort with potential, but requires "
        "refinement in gameplay clarity, audio balancing, and overall polish."
    )

    # -- McJames --
    doc.add_heading("3.2 Tester: McJames", level=2)

    add_table_from_data(
        doc,
        ["Category", "Rating / Feedback"],
        [
            ["Movement", "Works correctly"],
            ["Core Mechanics", "Understandable and functional"],
            ["Audio Design", "Feels unusual, though it fits the atmosphere"],
            ["Object Placement", "Collectable objects positioned too far apart"],
            ["Game Timer", "Could be increased to improve exploration and completion"],
        ],
    )
    doc.add_paragraph(
        "Summary: McJames found the core gameplay acceptable but noted that a longer "
        "timer and improved object placement would significantly improve gameplay flow."
    )

    # -- Gideon --
    doc.add_heading("3.3 Tester: Gideon", level=2)

    add_table_from_data(
        doc,
        ["Category", "Rating / Feedback"],
        [
            ["Movement", "Works correctly"],
            ["Recycling Mechanic", "Throwing objects into bins is enjoyable"],
            ["Jump Functionality", "Not working during testing"],
            ["Weather System", "Did not activate during testing"],
            ["Weather Potential", "Would have made gameplay more engaging if functional"],
        ],
    )
    doc.add_paragraph(
        "Summary: Gideon found the recycling mechanic enjoyable but identified two "
        "critical issues: the jump functionality not working and the weather system "
        "not activating. Debugging the weather system was identified as a priority."
    )

    # =====================================================================
    # 4. DEVELOPER TESTING
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("4. Developer Testing and Self Evaluation", level=1)

    doc.add_heading("4.1 Overall Design", level=2)
    doc.add_paragraph(
        "The overall design is acceptable but below the standard originally envisioned. "
        "Considering the available equipment and time constraints, this represents the "
        "best implementation achievable within the project period."
    )

    doc.add_heading("4.2 Development Challenges", level=2)
    add_table_from_data(
        doc,
        ["Challenge", "Impact"],
        [
            ["Limited lab access", "Entire development on personal laptop due to Unity version incompatibility"],
            ["Defining scope around SDG theme", "More difficult to scope appropriately than expected"],
            ["Early development", "Movement, jumping, and interaction progressed smoothly"],
            ["Environmental assets + weather integration", "Significantly increased project complexity"],
            ["UI and pause menu implementation", "Required additional prototyping scene and tooling"],
        ],
    )

    doc.add_heading("4.3 Weather System", level=2)
    doc.add_paragraph(
        "During early testing the weather system did not transition reliably. Investigation "
        "found this was an Inspector configuration issue (missing/mis-assigned references in "
        "WeatherFeedbackSystem). Following the 3 August fix pass the references were wired "
        "and proximity detection validated in batch mode."
    )
    doc.add_paragraph(
        "On 6 August the system was extended: wrong-bin recycling during a storm now "
        "physically pushes the player away from the bin (max 3 m/s, ramped), giving the "
        "weather state a direct gameplay consequence. Correct recycles calm the weather. "
        "The system was re-validated across three clean batch compile passes."
    )

    doc.add_heading("4.4 Reliable Systems", level=2)
    doc.add_paragraph(
        "The following systems have remained functional throughout development and across "
        "all testing sessions:"
    )
    add_bullet(doc, "Player movement — WASD controls, sprint, and jump physics")
    add_bullet(doc, "Mouse look — camera rotation, sensitivity, and invert-Y")
    add_bullet(doc, "Player interaction — pickup, drop, and throw mechanics")
    add_bullet(doc, "Object collection — item type identification and physics behaviour")
    add_bullet(doc, "Basic recycling mechanics — bin triggering and score updates")

    doc.add_heading("4.5 Project Organisation", level=2)
    doc.add_paragraph(
        "The folder structure and script architecture were reorganised multiple times "
        "to improve maintainability. The final structure organises scripts by system "
        "(Core, Player, Interaction, UI, Weather, Editor) with consistent naming "
        "conventions. This organisation makes project assets easier to locate and "
        "supports future development."
    )

    doc.add_heading("4.6 Assets Used", level=2)
    add_table_from_data(
        doc,
        ["Asset", "Usage", "Notes"],
        [
            ["Terrain URP Free Asset", "Terrain textures, soil, water plane", "Extensively used, suited the intended environment"],
            ["Rock Builder Free Asset", "Minor rock placement", "Used minimally; Terrain asset was more suitable"],
        ],
    )

    doc.add_heading("4.7 UI Improvements", level=2)
    doc.add_paragraph(
        "The Pause Menu and redesigned HUD were introduced to improve gameplay clarity "
        "and provide a cleaner user experience compared with the original interface. "
        "Time limitations prevented these systems from being refined further, particularly "
        "in terms of visual polish and transition animations."
    )

    doc.add_heading("4.8 Overall Reflection", level=2)
    doc.add_paragraph(
        "Although extensive testing and iteration were carried out, progress remained "
        "slower than anticipated. I am satisfied that I successfully produced a functional "
        "demonstration showcasing the core mechanics of the intended game. However, I am not "
        "satisfied with the overall quality of the final design, as it falls short of the "
        "original vision."
    )
    doc.add_paragraph(
        "Development began approximately one month before submission. Given additional "
        "development time, the project could have been significantly improved in terms of "
        "gameplay polish, environmental design, UI quality, and overall presentation."
    )
    doc.add_paragraph(
        "Following tutorials and researching different implementations, the weather system was "
        "successfully integrated so that it reacts dynamically to player recycling actions "
        "(correct recycles calm it, wrong recycles trigger a storm that pushes the player away). "
        "The most successful aspects of the project remain the implementation of the player "
        "movement and interaction systems, the recycling loop, and the dynamic weather feedback."
    )

    # =====================================================================
    # 5. TESTING SUMMARY
    # =====================================================================
    doc.add_page_break()
    doc.add_heading("5. Testing Summary", level=1)

    doc.add_heading("5.1 Successful Systems", level=2)
    add_bullet(doc, "Player movement (WASD, mouse look, sprint)")
    add_bullet(doc, "Object interaction (pickup, drop, throw)")
    add_bullet(doc, "Recycling mechanics and scoring")
    add_bullet(doc, "UI panel state machine (Welcome → Instructions → Playing → End)")
    add_bullet(doc, "TextMesh Pro HUD with event-driven updates")
    add_bullet(doc, "Input System action map switching")
    add_bullet(doc, "Audio ambient crossfade and SFX playback")
    add_bullet(doc, "Surface-based footstep audio")

    doc.add_heading("5.2 Partially Successful Systems", level=2)
    add_bullet(doc, "Weather system — fully wired on 3 Aug, extended with wind-push consequence on 6 Aug; validated in batch mode")
    add_bullet(doc, "Jump — ground check fixed on 6 Aug (CharacterController center + groundMask 81)")
    add_bullet(doc, "Pause menu — scripted and functional; requires Editor tool setup per scene")

    doc.add_heading("5.3 Known Bugs", level=2)
    add_bullet(doc, "8 toy instances + WaterPlane stacked off-world — needs repositioning in the Editor")
    add_bullet(doc, "Audio volume inconsistent between clips")
    add_bullet(doc, "Optimized SFX clips are untracked in git — must be delivered with the submission")

    doc.add_heading("5.4 Remaining Limitations", level=2)
    add_bullet(doc, "Collectable objects are spread too far for the 5-minute timer")
    add_bullet(doc, "No post-processing volume for visual polish")
    add_bullet(doc, "Pause menu panels are generated programmatically without styled backgrounds")
    add_bullet(doc, "Legacy RawAudio originals remain on disk (102 MB, untracked) as source provenance")

    doc.add_heading("5.5 Recommendations for Future Work", level=2)
    improvements = [
        "Reposition the 8 stacked toy instances and WaterPlane in the Editor",
        "Improve environmental readability and player guidance",
        "Balance audio levels across all clips",
        "Increase game timer or reduce object spread distance",
        "Add post-processing volume with bloom and tone mapping",
        "Implement full save/load system for game state persistence",
        "Refine UI animations and transition effects",
        "Perform broader playtesting with additional users",
    ]
    for imp in improvements:
        add_bullet(doc, imp)

    # =====================================================================
    # 6. CONCLUSION
    # =====================================================================
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "The Eco Rescue FPS project successfully demonstrates a functional first-person "
        "recycling simulation game built in Unity 6 with URP. The core gameplay loop — "
        "collect, sort, and recycle — is fully implemented and has been verified across "
        "multiple testing sessions."
    )
    doc.add_paragraph(
        "Player movement, interaction, and recycling mechanics are the project's strongest "
        "systems, providing a solid gameplay foundation. The UI has evolved from a basic "
        "immediate-start interface to a structured multi-panel system with welcome flow, "
        "pause functionality, settings configuration, and end-screen results."
    )
    doc.add_paragraph(
        "Peer testing identified three key areas requiring attention: weather system "
        "reliability, audio balancing, and gameplay pacing. The weather system has since "
        "been fully wired and extended with a wind-push consequence for wrong recycles; "
        "audio was remapped to a curated clip set; remaining issues are documented as "
        "known limitations with clear recommendations for future work."
    )
    doc.add_paragraph(
        "Despite the challenges encountered — including equipment limitations, software "
        "compatibility issues, and the complexity of integrating the dynamic weather system — "
        "the project meets the core requirements of the CM2121 assessment brief. The game "
        "is playable, the SDG 12 theme is clearly demonstrated through gameplay mechanics, "
        "and the scanned photogrammetry assets are properly integrated."
    )
    doc.add_paragraph(
        "The project is submitted with the understanding that it represents the best possible "
        "outcome given the available resources and timeframe, and that identified limitations "
        "have been fully documented for transparency and future improvement."
    )

    # -- Save --
    path = os.path.join(OUTPUT_DIR, "CM2121_Test_Log.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# =========================================================================
# MAIN
# =========================================================================

if __name__ == "__main__":
    print("Generating CM2121 documentation...")
    p1 = build_project_documentation()
    p2 = build_test_log()
    print(f"\nDone. Documents created:")
    print(f"  1. {p1}")
    print(f"  2. {p2}")
