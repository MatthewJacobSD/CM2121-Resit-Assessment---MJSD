#!/usr/bin/env python3
"""Convert DESIGN_DOCUMENT.md to a formatted .docx file."""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shading_elem)


def parse_inline(paragraph, text):
    """Parse bold and code inline formatting."""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        else:
            paragraph.add_run(part)


def parse_table(lines):
    """Parse markdown table lines into headers and rows."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if i == 0:
            headers = cells
        elif all(set(c.strip()) <= set('-: ') for c in cells):
            continue  # separator row
        else:
            rows.append(cells)
    return headers, rows


def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1', header))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        set_cell_shading(cell, '2E7D32')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(headers):
                cell = table.rows[i + 1].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                parse_inline(p, re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text))
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                if i % 2 == 0:
                    set_cell_shading(cell, 'F1F8E9')


def main():
    with open('docs/DESIGN_DOCUMENT.md', 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith('```') and not in_code_block:
            in_code_block = True
            code_lines = []
            i += 1
            continue
        elif line.strip().startswith('```') and in_code_block:
            in_code_block = False
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add light gray background via shading
            pPr = p._element.get_or_add_pPr()
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'F5F5F5',
            })
            pPr.append(shd)
            i += 1
            continue
        elif in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table handling
        if '|' in line and line.strip().startswith('|'):
            table_lines.append(line)
            i += 1
            continue
        elif table_lines:
            headers, rows = parse_table(table_lines)
            if headers and rows:
                add_table(doc, headers, rows)
            elif headers:
                add_table(doc, headers, [])
            table_lines = []

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('').paragraph_format.space_before = Pt(2)
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
            i += 1
            continue

        if line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
            i += 1
            continue

        if line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x43, 0x8A, 0x47)
            i += 1
            continue

        # Checkbox items
        if line.strip().startswith('- [x]'):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, line.strip()[6:])
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        if line.strip().startswith('- [ ]'):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, line.strip()[6:])
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # List items
        if line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, line.strip()[2:])
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Numbered list items
        m = re.match(r'^(\d+)\.\s+(.*)', line.strip())
        if m:
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, m.group(2))
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline(p, line.strip())
        p.paragraph_format.space_after = Pt(4)
        i += 1

    # Flush any remaining table
    if table_lines:
        headers, rows = parse_table(table_lines)
        if headers:
            add_table(doc, headers, rows)

    output_path = 'docs/CM2121_Design_Document.docx'
    doc.save(output_path)
    print(f'Saved: {output_path}')


if __name__ == '__main__':
    main()
